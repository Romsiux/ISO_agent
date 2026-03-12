"""
DOCX Exporter
=============
Converts a policy template (returned as Markdown-like text by the LLM)
into a properly formatted Word document (.docx) using python-docx.

Supports:
  - H1 / H2 / H3 headings (# / ## / ###)
  - Bold text (**text**)
  - Bullet lists (- item)
  - Numbered lists (1. item  or  - [ ] checkbox item)
  - Tables (| col | col |)
  - Horizontal rules (---)
  - Plain paragraphs
"""

from __future__ import annotations

import io
import re
from datetime import date
from typing import List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


# ── Colour palette ────────────────────────────────────────────────────────────
COLOUR_DARK_BLUE = RGBColor(0x1A, 0x29, 0x40)   # headings
COLOUR_MID_BLUE  = RGBColor(0x25, 0x63, 0xEB)   # H2
COLOUR_GREY      = RGBColor(0x6B, 0x72, 0x80)   # captions / hr


def _set_heading_colour(run, colour: RGBColor) -> None:
    run.font.color.rgb = colour


def _add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)


def _parse_inline(paragraph, text: str) -> None:
    """
    Parse inline **bold** and regular text and add runs to *paragraph*.
    """
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _parse_table(doc: Document, table_lines: List[str]) -> None:
    """
    Parse a Markdown table and add it to the document.

    table_lines  — the raw lines including header, separator, and data rows
    """
    # Filter out the separator row (---|---|---)
    rows = [l for l in table_lines if not re.match(r"^\|[\s\-|:]+\|$", l.strip())]
    if not rows:
        return

    # Parse cells
    parsed_rows = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed_rows.append(cells)

    if not parsed_rows:
        return

    col_count = max(len(r) for r in parsed_rows)
    table = doc.add_table(rows=len(parsed_rows), cols=col_count)
    table.style = "Table Grid"

    for r_idx, row_data in enumerate(parsed_rows):
        row_obj = table.rows[r_idx]
        for c_idx in range(col_count):
            cell = row_obj.cells[c_idx]
            text = row_data[c_idx] if c_idx < len(row_data) else ""
            # Strip bold markers — apply bold formatting
            is_bold = text.startswith("**") and text.endswith("**")
            clean_text = text.strip("*")
            p = cell.paragraphs[0]
            run = p.add_run(clean_text)
            if is_bold or r_idx == 0:
                run.bold = True
            p.paragraph_format.space_after = Pt(0)

    doc.add_paragraph()  # spacing after table


def policy_to_docx(policy_text: str, company_name: str = "Your Company") -> bytes:
    """
    Convert *policy_text* (Markdown-like) to a .docx byte stream.

    Returns raw bytes suitable for st.download_button.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Default font ──────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover strip ───────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover_run = cover.add_run(f"ISO 27001 / NIS2 Compliance Document")
    cover_run.font.color.rgb = COLOUR_GREY
    cover_run.font.size = Pt(9)
    cover.paragraph_format.space_after = Pt(2)

    company_p = doc.add_paragraph()
    cr = company_p.add_run(company_name)
    cr.bold = True
    cr.font.size = Pt(10)
    cr.font.color.rgb = COLOUR_DARK_BLUE
    company_p.paragraph_format.space_after = Pt(2)

    date_p = doc.add_paragraph()
    date_run = date_p.add_run(f"Generated: {date.today().strftime('%d %B %Y')}")
    date_run.font.color.rgb = COLOUR_GREY
    date_p.paragraph_format.space_after = Pt(12)

    _add_horizontal_rule(doc)

    # ── Parse body ────────────────────────────────────────────────────────────
    lines = policy_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Table — collect all consecutive table lines
        if stripped.startswith("|"):
            table_block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_block.append(lines[i])
                i += 1
            _parse_table(doc, table_block)
            continue

        # Headings
        h3 = re.match(r"^### (.+)$", stripped)
        h2 = re.match(r"^## (.+)$", stripped)
        h1 = re.match(r"^# (.+)$", stripped)

        if h1:
            p = doc.add_heading(h1.group(1), level=1)
            for run in p.runs:
                run.font.color.rgb = COLOUR_DARK_BLUE
                run.font.size = Pt(18)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        if h2:
            p = doc.add_heading(h2.group(1), level=2)
            for run in p.runs:
                run.font.color.rgb = COLOUR_MID_BLUE
                run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        if h3:
            p = doc.add_heading(h3.group(1), level=3)
            for run in p.runs:
                run.font.color.rgb = COLOUR_DARK_BLUE
                run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Checkbox bullet  - [ ]
        checkbox = re.match(r"^- \[[ x]\] (.+)$", stripped)
        if checkbox:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25)
            tick = "☑ " if "x" in stripped[3] else "☐ "
            run = p.add_run(tick)
            run.font.size = Pt(11)
            _parse_inline(p, checkbox.group(1))
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Bullet list  - item
        bullet = re.match(r"^[-*] (.+)$", stripped)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            _parse_inline(p, bullet.group(1))
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Numbered list  1. item
        numbered = re.match(r"^\d+\. (.+)$", stripped)
        if numbered:
            p = doc.add_paragraph(style="List Number")
            _parse_inline(p, numbered.group(1))
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Plain paragraph
        p = doc.add_paragraph()
        _parse_inline(p, stripped)
        p.paragraph_format.space_after = Pt(6)
        i += 1

    # ── Footer ────────────────────────────────────────────────────────────────
    _add_horizontal_rule(doc)
    footer_p = doc.add_paragraph()
    footer_r = footer_p.add_run(
        f"Generated by ISO 27001 / NIS2 Compliance Assistant  |  {company_name}  |  CONFIDENTIAL"
    )
    footer_r.font.size = Pt(8)
    footer_r.font.color.rgb = COLOUR_GREY
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()